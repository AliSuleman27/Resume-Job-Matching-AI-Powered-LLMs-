import os
import re
import json
import logging
from groq import Groq
from PyPDF2 import PdfReader
from docx import Document
from flask import current_app
from prompt import generate_resume_prompt, generate_job_prompt

logger = logging.getLogger(__name__)


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in current_app.config['ALLOWED_EXTENSIONS']


def extract_json_from_response(content: str) -> str:
    """Extracts JSON content from the LLM response."""
    json_match = re.search(r'```json\n(.*?)\n```', content, re.DOTALL)
    if json_match:
        return json_match.group(1).strip()

    try:
        start = content.find('{')
        end = content.rfind('}') + 1
        if start != -1 and end != -1:
            potential_json = content[start:end]
            json.loads(potential_json)
            return potential_json
    except json.JSONDecodeError:
        pass

    return content.strip()


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extracts text from different file types with encoding fallback."""
    try:
        if file_type == 'pdf':
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                text = '\n'.join([page.extract_text() for page in reader.pages])
        elif file_type == 'docx':
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
        else:
            encodings = ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
        return text
    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        raise


def extract_text_from_bytes(raw: bytes, file_type: str) -> str:
    """Extract text from raw file bytes (works on any platform including Vercel)."""
    import io

    if not raw:
        raise ValueError("Uploaded file is empty")

    try:
        if file_type == 'pdf':
            reader = PdfReader(io.BytesIO(raw))
            text = '\n'.join([page.extract_text() or '' for page in reader.pages])
        elif file_type == 'docx':
            doc = Document(io.BytesIO(raw))
            text = '\n'.join([para.text for para in doc.paragraphs])
        else:
            for encoding in ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1']:
                try:
                    text = raw.decode(encoding)
                    break
                except (UnicodeDecodeError, AttributeError):
                    continue
            else:
                text = raw.decode('utf-8', errors='ignore')
        return text
    except Exception as e:
        logger.error(f"Error extracting text from bytes: {e}")
        raise


def extract_text_from_stream(file_storage, file_type: str) -> str:
    """Extract text directly from a Werkzeug FileStorage stream (no disk I/O).

    This avoids filesystem issues on serverless platforms like Vercel where
    saving and re-reading files can corrupt binary data.
    """
    import io

    try:
        file_storage.stream.seek(0)
        raw = file_storage.stream.read()

        if not raw:
            raise ValueError("Uploaded file is empty")

        if file_type == 'pdf':
            reader = PdfReader(io.BytesIO(raw))
            text = '\n'.join([page.extract_text() or '' for page in reader.pages])
        elif file_type == 'docx':
            doc = Document(io.BytesIO(raw))
            text = '\n'.join([para.text for para in doc.paragraphs])
        else:
            for encoding in ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1']:
                try:
                    text = raw.decode(encoding)
                    break
                except (UnicodeDecodeError, AttributeError):
                    continue
            else:
                text = raw.decode('utf-8', errors='ignore')
        return text
    except Exception as e:
        logger.error(f"Error extracting text from stream: {e}")
        raise


def call_llm(resume_text: str) -> dict:
    """Calls the Groq API to parse resume text."""
    try:
        client = Groq(api_key=os.getenv('GROQ_API_KEY'))
        prompt = generate_resume_prompt(resume_text)

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a helpful resume parser."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        content = response.choices[0].message.content
        tokens = response.usage.total_tokens if hasattr(response, "usage") else 0

        json_content = extract_json_from_response(content)
        return {"output": json_content, "tokens": tokens}
    except Exception as e:
        logger.error(f"Error in call_llm: {str(e)}")
        return {"output": "", "error": str(e), "tokens": 0}


def call_job_llm(job_text: str) -> dict:
    """Calls the Groq API to parse job description text."""
    try:
        client = Groq(api_key=os.getenv('GROQ_API_KEY'))
        prompt = generate_job_prompt(job_text)

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a helpful job description parser."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        content = response.choices[0].message.content
        tokens = response.usage.total_tokens if hasattr(response, "usage") else 0

        json_content = extract_json_from_response(content)
        return {"output": json_content, "tokens": tokens}
    except Exception as e:
        logger.error(f"Error in call_job_llm: {str(e)}")
        return {"output": "", "error": str(e), "tokens": 0}
