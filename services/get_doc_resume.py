import json
from pydantic import ValidationError
from typing import Any
from models.resume_model import Resume
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
from pydantic import ValidationError
from typing import Dict

def parse_resume_from_dict(python_dict: Dict) -> Resume:
    """
    Converts a Python dictionary into JSON and parses it into a Resume Pydantic model.

    Args:
        python_dict (Dict): Python dictionary containing resume data.

    Returns:
        Resume: Pydantic Resume object.

    Raises:
        ValidationError: If dictionary structure does not match the Resume model.
    """
    try:
        json_str = json.dumps(python_dict)  # Convert dict to JSON string
        json_data = json.loads(json_str)    # Parse JSON string back to dict (can be skipped, here for clarity)
        resume_obj = Resume(**json_data)    # Create Resume model
        return resume_obj
    except ValidationError as e:
        print("Validation error:", e)
        raise

def create_pretty_resume_docx(resume: Resume, output_path: str) -> None:
    """
    Generate a pretty resume document (.docx) from a Resume Pydantic object.

    Args:
        resume (Resume): The resume object to convert.
        output_path (str): File path to save the generated .docx document.
    """
    doc = Document()

    # Title - Full Name
    title = doc.add_heading(resume.basic_info.full_name, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Current Title
    if resume.basic_info.current_title:
        current_title = doc.add_paragraph(resume.basic_info.current_title)
        current_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact Info (Email, Phone, LinkedIn, GitHub, Portfolio, Address)
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    contacts = []
    if resume.contact_info.email:
        contacts.append(f"Email: {resume.contact_info.email}")
    if resume.contact_info.phone:
        contacts.append(f"Phone: {resume.contact_info.phone}")
    if resume.contact_info.linkedin:
        contacts.append(f"LinkedIn: {resume.contact_info.linkedin}")
    if resume.contact_info.github:
        contacts.append(f"GitHub: {resume.contact_info.github}")
    if resume.contact_info.portfolio:
        contacts.append(f"Portfolio: {resume.contact_info.portfolio}")

    # Format address if available
    addr = resume.contact_info.address
    if addr:
        parts = [addr.street, addr.city, addr.state, addr.zip_code, addr.country]
        addr_str = ", ".join(filter(None, parts))
        if addr_str:
            contacts.append(f"Address: {addr_str}")

    contact_info.add_run(" | ".join(contacts)).bold = True
    doc.add_paragraph()  # spacer

    # Summary (optional)
    if resume.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume.summary)
        doc.add_paragraph()

    # Skills grouped by category, include proficiency and years_of_experience if available
    if resume.skills:
        doc.add_heading("Skills", level=1)
        skill_lines = {}
        for skill in resume.skills:
            cat = skill.category or "Other"
            skill_text = skill.skill_name
            details = []
            if skill.proficiency:
                details.append(skill.proficiency)
            if skill.years_of_experience is not None:
                details.append(f"{skill.years_of_experience} yrs")
            if details:
                skill_text += f" ({', '.join(details)})"
            skill_lines.setdefault(cat, []).append(skill_text)

        for cat, skills_list in skill_lines.items():
            doc.add_heading(cat, level=2)
            doc.add_paragraph(", ".join(skills_list))
        doc.add_paragraph()

    # Education
    if resume.education:
        doc.add_heading("Education", level=1)
        for edu in resume.education:
            p = doc.add_paragraph()
            # Degree, Field, Institution
            degree_field = edu.degree
            if edu.field:
                degree_field += f" in {edu.field}"
            inst = edu.institution or "N/A"
            p.add_run(f"{degree_field} - {inst}").bold = True

            # Dates and location
            dates = ""
            if edu.start_date or edu.end_date:
                dates = f" ({edu.start_date or '?'} - {edu.end_date or '?'})"
            loc = f", {edu.location}" if edu.location else ""
            p.add_run(dates + loc).italic = True

            # Grade if present
            if edu.grade:
                p.add_run(f" | Grade: {edu.grade}").italic = True

            # Courses list if any
            if edu.courses:
                courses = ", ".join(edu.courses)
                doc.add_paragraph(f"Courses: {courses}", style='List Bullet')
        doc.add_paragraph()

    # Experience
    if resume.experience:
        doc.add_heading("Experience", level=1)
        for exp in resume.experience:
            p = doc.add_paragraph()
            p.add_run(f"{exp.job_title} at {exp.company or 'N/A'}").bold = True

            dates = f"{exp.start_date or '?'} - {exp.end_date or 'Present'}"
            p.add_run(f" ({dates})").italic = True

            if exp.location:
                p.add_run(f", {exp.location}")
            if exp.employment_type:
                p.add_run(f" | {exp.employment_type}").italic = True

            # Responsibilities
            for resp in exp.responsibilities:
                doc.add_paragraph(resp, style='List Bullet')

            # Skills used in the experience
            if exp.skills_used:
                doc.add_paragraph(f"Skills Used: {', '.join(exp.skills_used)}", style='List Bullet')

            doc.add_paragraph()

    # Projects
    if resume.projects:
        doc.add_heading("Projects", level=1)
        for proj in resume.projects:
            p = doc.add_paragraph()
            p.add_run(proj.title).bold = True

            dates = ""
            if proj.start_date or proj.end_date:
                dates = f" ({proj.start_date or '?'} - {proj.end_date or '?'})"
            p.add_run(dates).italic = True

            if proj.role:
                p.add_run(f" | Role: {proj.role}").italic = True

            if proj.link:
                p.add_run(f" [Link]").italic = True

            if proj.description:
                doc.add_paragraph(proj.description)

            if proj.technologies:
                doc.add_paragraph(f"Technologies: {', '.join(proj.technologies)}", style='List Bullet')
            doc.add_paragraph()

    # Certifications
    if resume.certifications:
        doc.add_heading("Certifications", level=1)
        for cert in resume.certifications:
            p = doc.add_paragraph()
            p.add_run(cert.title).bold = True

            if cert.issuer:
                p.add_run(f" - {cert.issuer}")

            dates = ""
            if cert.issue_date or cert.expiration_date:
                dates = f" ({cert.issue_date or '?'} - {cert.expiration_date or 'N/A'})"
            p.add_run(dates).italic = True

            if cert.credential_id:
                p.add_run(f" | Credential ID: {cert.credential_id}").italic = True

            if cert.url:
                p.add_run(f" [Link]").italic = True
        doc.add_paragraph()

    # Publications (if any)
    if resume.publications:
        doc.add_heading("Publications", level=1)
        for pub in resume.publications:
            p = doc.add_paragraph()
            p.add_run(pub.title).bold = True
            details = []
            if pub.journal:
                details.append(pub.journal)
            if pub.year:
                details.append(str(pub.year))
            if details:
                p.add_run(f" ({', '.join(details)})").italic = True
            if pub.authors:
                p.add_paragraph(f"Authors: {', '.join(pub.authors)}", style='List Bullet')
            if pub.doi:
                p.add_paragraph(f"DOI: {pub.doi}", style='List Bullet')
            if pub.url:
                p.add_run(f" [Link]").italic = True
        doc.add_paragraph()

    # Languages
    if resume.languages:
        doc.add_heading("Languages", level=1)
        for lang in resume.languages:
            line = lang.language
            if lang.proficiency:
                line += f" ({lang.proficiency})"
            doc.add_paragraph(line, style='List Bullet')
        doc.add_paragraph()

    # Interests
    if resume.interests:
        doc.add_heading("Interests", level=1)
        for interest in resume.interests:
            doc.add_paragraph(interest, style='List Bullet')
        doc.add_paragraph()

    # Volunteer Experience
    if resume.volunteer_experience:
        doc.add_heading("Volunteer Experience", level=1)
        for vol in resume.volunteer_experience:
            p = doc.add_paragraph()
            p.add_run(vol.role).bold = True
            if vol.organization:
                p.add_run(f" at {vol.organization}")
            dates = ""
            if vol.start_date or vol.end_date:
                dates = f" ({vol.start_date or '?'} - {vol.end_date or 'Present'})"
            p.add_run(dates).italic = True
            if vol.description:
                doc.add_paragraph(vol.description)
            doc.add_paragraph()

    # Awards
    if resume.awards:
        doc.add_heading("Awards", level=1)
        for award in resume.awards:
            line = award.title
            if award.issuer:
                line += f" - {award.issuer}"
            if award.date:
                line += f" ({award.date})"
            doc.add_paragraph(line, style='List Bullet')
            if award.description:
                doc.add_paragraph(award.description)
        doc.add_paragraph()

    # Links (if any)
    if resume.links:
        doc.add_heading("Links", level=1)
        for link in resume.links:
            label = link.label or "Link"
            url = link.url or ""
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(url).italic = True
        doc.add_paragraph()

    # Save the document
    doc.save(output_path)


