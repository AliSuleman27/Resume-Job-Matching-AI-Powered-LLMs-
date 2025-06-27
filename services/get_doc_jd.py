import json
from pydantic import ValidationError
from models.job_description_model import JobDescription
from docx import Document
from typing import Optional
import json
from pydantic import ValidationError
from docx.shared import Pt

def parse_job_description_from_file(python_dict: dict) -> JobDescription:
    """
    Reads a JSON file from the given path and parses it into a JobDescription Pydantic model.

    Args:
        file_path (str): Path to the JSON file.

    Returns:
        JobDescription: Pydantic JobDescription object.

    Raises:
        FileNotFoundError: If the file does not exist.
        json.JSONDecodeError: If the file content is not valid JSON.
        ValidationError: If JSON structure does not match the JobDescription model.
    """
    try:
        json_str = json.dumps(python_dict)  # Convert dict to JSON string
        json_data = json.loads(json_str)    # Parse JSON string back to dict (can be skipped, here for clarity)
        jd_object = JobDescription(**json_data)    # Create Resume model
        return jd_object
    except ValidationError as e:
        print("Validation error:", e)
        raise

def create_pretty_jd_docx(jd: JobDescription, output_path: str) -> None:
    doc = Document()

    def add_heading(text: str, level: int = 1):
        doc.add_heading(text, level=level)

    def add_paragraph(text: str, bold: bool = False, italic: bool = False):
        if not text:
            return
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        p.paragraph_format.space_after = Pt(6)

    def add_list(items: Optional[list[str]], title: Optional[str] = None):
        if items:
            if title:
                add_heading(title, level=2)
            for item in items:
                doc.add_paragraph(item, style='List Bullet')

    # Title and Company
    add_heading(jd.title, level=0)
    if jd.company and jd.company.name:
        add_paragraph(f"Company: {jd.company.name}", bold=True)

    # Summary
    if jd.summary:
        add_heading("Summary", level=2)
        add_paragraph(jd.summary)

    # Description
    if jd.description:
        add_heading("Job Description", level=2)
        add_paragraph(jd.description)

    # Employment & Job Info
    add_heading("Job Details", level=2)
    add_paragraph(f"Employment Type: {jd.employment_type}")
    add_paragraph(f"Job Level: {jd.job_level}")
    add_paragraph(f"Industry: {jd.industry}")
    add_paragraph(f"Function: {jd.function}")
    add_paragraph(f"Department: {jd.department}")
    add_paragraph(f"Remote: {jd.is_remote}, Hybrid: {jd.is_hybrid}, Onsite: {jd.is_onsite}")

    # Location(s)
    if jd.locations:
        add_heading("Locations", level=2)
        for loc in jd.locations:
            loc_text = ", ".join(filter(None, [loc.city, loc.state, loc.country, loc.zip_code]))
            if loc.remote:
                loc_text += " (Remote)"
            add_paragraph(loc_text)

    # Salary
    if jd.salary:
        salary = jd.salary
        sal_text = f"{salary.currency or ''} {salary.min or ''} - {salary.max or ''} per {salary.period or ''}"
        if salary.is_estimated:
            sal_text += " (Estimated)"
        add_heading("Salary", level=2)
        add_paragraph(sal_text)

    # Qualifications
    if jd.qualifications:
        add_heading("Qualifications", level=2)
        if jd.qualifications.education:
            for edu in jd.qualifications.education:
                add_paragraph(f"{edu.level or ''} in {edu.field_of_study or ''}")
        if jd.qualifications.experience_years:
            exp = jd.qualifications.experience_years
            add_paragraph(f"Experience: {exp.min} - {exp.max} years")
        if jd.qualifications.certifications:
            add_list(jd.qualifications.certifications, title="Certifications")

    # Skills
    if jd.skills:
        add_heading("Skills", level=2)
        add_list(jd.skills.mandatory, "Mandatory")
        add_list(jd.skills.optional, "Optional")
        add_list(jd.skills.tools, "Tools")

    # Languages
    if jd.languages:
        add_heading("Language Requirements", level=2)
        for lang in jd.languages:
            add_paragraph(f"{lang.language} - {lang.proficiency}")

    # Responsibilities & Requirements
    add_list(jd.responsibilities, title="Responsibilities")
    add_list(jd.requirements, title="Requirements")
    add_list(jd.nice_to_have, title="Nice to Have")

    # Benefits
    add_list(jd.benefits, title="Benefits")

    # Application Info
    if jd.application_url:
        add_paragraph(f"Apply here: {jd.application_url}")

    # Posting Info
    add_paragraph(f"Posted on: {jd.posting_date}")
    add_paragraph(f"Closing on: {jd.closing_date}")

    # Save Document
    doc.save(output_path)
